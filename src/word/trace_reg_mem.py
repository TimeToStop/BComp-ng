from docx.shared import Cm
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK


#
#   Set text to table cell with:
#       style
#       vertical, horizontal align
#       if width = -1 => auto else width
#


def add_text(table, i, j, style, width, text):
    table.cell(i, j).paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(i, j).paragraphs[-1].style = style
    table.cell(i, j).paragraphs[-1].text = text
    if width != -1:
        table.cell(i, j).width = Cm(width)
    table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table_header(table, style):
    print('Making table header.')
    table.cell(0, 0).merge(table.cell(0, 1))
    add_text(table, 0, 0, style, -1, 'Выполняемая команда')

    table.cell(0, 2).merge(table.cell(0, 3))
    table.cell(0, 3).merge(table.cell(0, 4))
    table.cell(0, 4).merge(table.cell(0, 5))
    table.cell(0, 5).merge(table.cell(0, 6))
    table.cell(0, 6).merge(table.cell(0, 7))
    table.cell(0, 7).merge(table.cell(0, 8))
    table.cell(0, 8).merge(table.cell(0, 9))

    add_text(table, 0, 2, style, -1, 'Содержимое регистров процессора после выполнения команды')

    table.cell(0, 10).merge(table.cell(0, 11))
    add_text(table, 0, 10, style, -1, 'Ячейка, содержимое которой изменилось после выполнения команды')

    head = [
        'Адрес',
        'Код',
        'IP',
        'CR',
        'AR',
        'DR',
        'SP',
        'BR',
        'AC',
        'NZVC',
        'Адрес',
        'Новый код'
    ]
    i = 0

    for d in head:
        add_text(table, 1, i, style, -1, d)
        i += 1


def trace(regs, diff_mem, start_program, d):
    # Add set vertical margin very small
    section = d.sections[-1]
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    # Add table style
    style = d.styles.add_style('Trace Table Style', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.italic = True
    style.font.bold = True

    # Add table
    table = d.add_table(len(regs) + 2, len(regs[0]) + 4)
    table.style = 'Table Grid'
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Generate first 2 lines
    table_header(table, style)

    # Fill table
    print('Fill table:')
    # Current IP is previous IP so we hold previous IP
    # First Previous IP is program start
    previous_ip = start_program[1:-1:]

    for i in range(len(regs)):
        print('\tFill line {}/{}.'.format(i + 1, len(regs)))

        # Address of cell that was executed
        add_text(table, i + 2, 0, style, 2, previous_ip)
        # Command that was executed(same value than in Command Register(1))
        add_text(table, i + 2, 1, style, 2, regs[i][1])
        # Previous IP for next command is current IP(0)
        previous_ip = regs[i][0]

        # Register value
        for j in range(len(regs[0])):
            add_text(table, i + 2, j + 2, style, 1, regs[i][j])

        # Addresses and values that had changed
        addr_changed = ''
        val_changed = ''

        for diff in diff_mem[i]:
            addr_changed += diff[0] + '\n'
            val_changed += diff[1] + '\n'

        add_text(table, i + 2, len(regs[0]) + 2, style, 2, addr_changed[:-1:])
        add_text(table, i + 2, len(regs[0]) + 3, style, 2, val_changed[:-1:])

    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
