from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from .trace_reg_mem import add_text
from src.comments.comments import comment

def init_header(table, style):
    header = ['№', 'Адрес', 'Команды', 'Мнемоника', 'Комментарий']

    add_text(table, 0, 0, style, 0.1, header[0])
    add_text(table, 0, 1, style, 2, header[1])
    add_text(table, 0, 2, style, 2, header[2])
    add_text(table, 0, 3, style, 5, header[3])
    add_text(table, 0, 4, style, 5, header[4])


def create_func(func, style, d):
    table = d.add_table(len(func), 5)
    table.style = 'Table Grid'
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    init_header(table, style)
    ip = int(func[0], 16)

    for i in range(1, len(func)):
        print('\t\tCommand {}/{}:'.format(i, len(func) - 1))
        mnem, comm = comment(ip + 1, func[i])
        add_text(table, i, 0, style, 0.1, str(i))
        add_text(table, i, 1, style, 2, hex(ip)[2::].zfill(3).upper())
        add_text(table, i, 2, style, 2, func[i])
        add_text(table, i, 3, style, 5, mnem)
        add_text(table, i, 4, style, 5, comm)
        ip += 1

    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    

def program_table(program, d):
    # Add table style
    style = d.styles.add_style('Program Table Style', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.italic = True

    i = 1
    print('\n\nFill program table:')

    for func in program:
        print('\tFunction {}/{}:'.format(i, len(program)))
        create_func(func, style, d)
        i += 1
